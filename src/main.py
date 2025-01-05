import base64
import uuid
from contextlib import asynccontextmanager
from io import BytesIO
from operator import or_
from uuid import UUID

from docx import Document
from docx.shared import Inches
from fastapi import Depends, FastAPI
from fastapi.responses import StreamingResponse
from fastapi.security import OAuth2PasswordBearer
from sqlalchemy.orm import Session

from .auth import verify_access_token
from .database import create_tables, get_db
from .models import Chat, Message


@asynccontextmanager
async def lifespan(_: FastAPI):
    # Startup code
    create_tables()
    yield


app = FastAPI(lifespan=lifespan)

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")

role_mapping = {
    "user": "Vi",
    "assistant": "Asistent",
}


@app.get("/export-document/{chat_id}")
async def export_document(
    chat_id: UUID,
    db: Session = Depends(get_db),
    token: str = Depends(oauth2_scheme),
):
    payload = verify_access_token(token)
    if not payload:
        return {"msg": "Invalid token"}
    user_id = payload["sub"]
    chat = db.query(Chat).filter(Chat.id == chat_id, Chat.user_id == user_id).first()
    if not chat:
        return {"msg": "Chat not found"}

    previous_messages = (
        db.query(Message)
        .filter(
            Message.chat_id == chat_id,
            or_(Message.role == "user", Message.role == "assistant"),
        )
        .all()
    )

    doc = Document()
    doc.add_heading(chat.name, level=1)

    for message in previous_messages:
        paragraph = doc.add_paragraph()
        # role = (role_mapping[message.role] or "")
        role_run = paragraph.add_run(
            f"{(role_mapping[message.role] or "").capitalize()}: "
        )
        role_run.bold = True  # Make the role bold

        # Add text content if present
        if message.content:
            paragraph.add_run(message.content)

        # Check for an image field in the message
        if message.base64_image != "":
            try:
                # Strip the prefix and decode the base64 string
                _, encoded_image = message.base64_image.split(",", 1)
                image_data = base64.b64decode(encoded_image)

                # Add a new paragraph for the image
                image_stream = BytesIO(image_data)
                doc.add_picture(image_stream, width=Inches(5))  # Adjust size as needed
            except Exception as e:
                doc.add_paragraph(f"[Failed to render image: {str(e)}]")

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"{chat.name.replace(' ', '_')}_{uuid.uuid4()}.docx"

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
